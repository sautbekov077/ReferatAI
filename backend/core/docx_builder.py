from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, uuid
from backend.models.schemas import ExportDocxRequest, Section

def build_docx(req: ExportDocxRequest) -> str:
    doc = Document()
    section = doc.sections[0]
    margin = Cm(req.margins_cm)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = margin

    style = doc.styles['Normal']
    style.font.name = req.font_name
    style.font.size = Pt(req.font_size_pt)

    # Титульный — по центру
    title = doc.add_paragraph(req.topic)
    title.style = doc.styles['Title']
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Оглавление-напоминание
    if req.include_toc:
        h_toc = doc.add_paragraph("Оглавление")
        h_toc.style = 'Heading 1'  # заголовки обычно выровнены влево
        tip = doc.add_paragraph("Сформировать в Word: Ссылки → Оглавление (таблица содержимого).")
        tip.italic = True
        tip.paragraph_format.line_spacing = req.line_spacing
        tip.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # поясняющий текст тоже по ширине
        doc.add_page_break()

    def _is_refs(title: str) -> bool:
        t = (title or "").lower()
        return "список литературы" in t or "references" in t or "список источников" in t

    def _apply_hanging_indent(paragraph, cm=1.0):
        pf = paragraph.paragraph_format
        pf.left_indent = Cm(cm)
        pf.first_line_indent = Cm(-cm)
        pf.line_spacing = req.line_spacing
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # источники также по ширине [вид аккуратнее]

    def write_section(s: Section, level: int = 1):
        # Заголовок раздела — стандартный стиль (обычно слева)
        h = doc.add_paragraph(s.title)
        h.style = f"Heading {min(level,3)}"

        refs_mode = _is_refs(s.title)
        for par in s.paragraphs:
            p = doc.add_paragraph(par)
            p.paragraph_format.line_spacing = req.line_spacing
            # Ключевое: выравнивание основного текста по ширине
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # выровнять по ширине [web:46][web:32][web:41]
            if refs_mode:
                _apply_hanging_indent(p, cm=1.0)

        for sub in s.subsections:
            write_section(sub, level+1)

    write_section(req.essay, 1)

    outdir = "output"
    os.makedirs(outdir, exist_ok=True)
    filename = os.path.join(outdir, f"{req.doc_type}_{uuid.uuid4().hex}.docx")
    doc.save(filename)
    return filename
